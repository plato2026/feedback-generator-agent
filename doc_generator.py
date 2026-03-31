"""
doc_generator.py — Word Document Generator
===========================================
Creates professionally formatted .docx feedback documents
using python-docx. Generates proper Word files (not HTML-as-doc).
"""

import io
import re
from datetime import datetime
from typing import List, Dict, Optional

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_ORIENT


# ─── COLOR CONSTANTS ──────────────────────────────────────────────

DARK_BLUE = RGBColor(0x1B, 0x3A, 0x5C)
MED_BLUE = RGBColor(0x2E, 0x5C, 0x8A)
LIGHT_BLUE = RGBColor(0x3D, 0x7A, 0xB5)
PURPLE = RGBColor(0x5A, 0x3E, 0x7A)
DARK_TEXT = RGBColor(0x33, 0x33, 0x33)
MUTED_TEXT = RGBColor(0x66, 0x66, 0x66)
ACCENT = RGBColor(0xC0, 0x60, 0x30)


def _setup_styles(doc: Document):
    """Configure document styles for consistent formatting."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Georgia"
    font.size = Pt(11)
    font.color.rgb = DARK_TEXT
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # Heading 1
    h1 = doc.styles["Heading 1"]
    h1.font.name = "Georgia"
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = DARK_BLUE
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)

    # Heading 2
    h2 = doc.styles["Heading 2"]
    h2.font.name = "Georgia"
    h2.font.size = Pt(15)
    h2.font.bold = True
    h2.font.color.rgb = MED_BLUE
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(8)

    # Heading 3
    h3 = doc.styles["Heading 3"]
    h3.font.name = "Georgia"
    h3.font.size = Pt(13)
    h3.font.bold = True
    h3.font.color.rgb = LIGHT_BLUE
    h3.paragraph_format.space_before = Pt(14)
    h3.paragraph_format.space_after = Pt(6)


def _add_title_page(doc: Document, category_label: str, date_str: str):
    """Add a formatted title page."""
    # Spacer
    for _ in range(6):
        doc.add_paragraph("")

    # Category label
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(category_label.upper())
    run.font.name = "Georgia"
    run.font.size = Pt(10)
    run.font.color.rgb = MUTED_TEXT
    run.font.letter_spacing = Pt(3)

    # Main title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Feedback Document")
    run.font.name = "Georgia"
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = DARK_BLUE

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Comprehensive Assessment & Expert Synthesis")
    run.font.name = "Georgia"
    run.font.size = Pt(14)
    run.font.color.rgb = MED_BLUE

    # Spacer
    doc.add_paragraph("")
    doc.add_paragraph("")

    # Institution info
    meta_lines = [
        "Human-AI Collaborative Systems",
        "Luddy School of Informatics, Computing, and Engineering",
        "Indiana University Indianapolis",
        "",
        "Professor Fawzi BenMessaoud",
        date_str,
    ]
    for line in meta_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.name = "Georgia"
        run.font.size = Pt(11)
        run.font.color.rgb = MUTED_TEXT

    # Page break
    doc.add_page_break()


def _add_part_label(doc: Document, label: str):
    """Add a small uppercase part label (e.g., 'PART A')."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(label)
    run.font.name = "Georgia"
    run.font.size = Pt(10)
    run.font.color.rgb = MUTED_TEXT
    run.font.letter_spacing = Pt(3)
    run.bold = True


def _add_markdown_content(doc: Document, markdown_text: str):
    """
    Parse simple markdown formatting and add it to the document.
    Handles **bold**, *italic*, and ### headings.
    """
    if not markdown_text:
        return

    paragraphs = markdown_text.split("\n\n")
    for para_text in paragraphs:
        para_text = para_text.strip()
        if not para_text:
            continue

        # Check for heading patterns
        if para_text.startswith("### "):
            doc.add_heading(para_text[4:].strip(), level=3)
            continue
        elif para_text.startswith("## "):
            doc.add_heading(para_text[3:].strip(), level=2)
            continue
        elif para_text.startswith("# "):
            doc.add_heading(para_text[2:].strip(), level=1)
            continue

        # Handle bold markdown headers like **Section A1: ...**
        if para_text.startswith("**") and para_text.endswith("**") and "\n" not in para_text:
            clean = para_text.strip("*").strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            run = p.add_run(clean)
            run.font.name = "Georgia"
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = PURPLE
            continue

        # Regular paragraph with inline formatting
        p = doc.add_paragraph()
        _add_formatted_runs(p, para_text)


def _add_formatted_runs(paragraph, text: str):
    """Add text to a paragraph with **bold** and *italic* formatting."""
    # Split by bold and italic markers
    # Pattern: find **bold**, *italic*, or plain text
    pattern = r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*)'
    parts = re.split(pattern, text)

    for part in parts:
        if not part:
            continue

        if part.startswith("***") and part.endswith("***"):
            clean = part[3:-3]
            run = paragraph.add_run(clean)
            run.bold = True
            run.italic = True
        elif part.startswith("**") and part.endswith("**"):
            clean = part[2:-2]
            run = paragraph.add_run(clean)
            run.bold = True
            run.font.color.rgb = MED_BLUE
        elif part.startswith("*") and part.endswith("*"):
            clean = part[1:-1]
            run = paragraph.add_run(clean)
            run.italic = True
        else:
            # Handle line breaks within paragraphs
            lines = part.split("\n")
            for i, line in enumerate(lines):
                if i > 0:
                    paragraph.add_run("\n")
                run = paragraph.add_run(line)
                run.font.name = "Georgia"
                run.font.size = Pt(11)


def _add_student_section(doc: Document, name: str, feedback: str):
    """Add a single student's feedback section with visual formatting."""
    # Student name as heading
    doc.add_heading(name, level=2)

    # Add feedback content
    _add_markdown_content(doc, feedback)

    # Divider line
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    # Add a thin horizontal rule via bottom border
    pf = p.paragraph_format
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    p._p.get_or_add_pPr().append(pBdr)


def generate_feedback_docx(
    category_label: str,
    part_a_label: Optional[str],
    part_b_label: str,
    has_part_a: bool,
    part_a_results: List[Dict[str, str]],
    part_b_text: str,
) -> bytes:
    """
    Generate a complete feedback .docx document.

    Args:
        category_label: e.g., "Hands-On Learning Activity"
        part_a_label: e.g., "Individualized Student Feedback"
        part_b_label: e.g., "Collective Expert Synthesis — Masterclass"
        has_part_a: whether Part A (individual feedback) exists
        part_a_results: list of {"name": str, "feedback": str}
        part_b_text: the collective synthesis text

    Returns:
        Bytes of the generated .docx file.
    """
    doc = Document()

    # Page setup: US Letter, 1-inch margins
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Apply styles
    _setup_styles(doc)

    # Title page
    date_str = datetime.now().strftime("%B %d, %Y")
    _add_title_page(doc, category_label, date_str)

    # ─── Part A: Individual Feedback ──────────────────────────────
    if has_part_a and part_a_results:
        _add_part_label(doc, "PART A")
        doc.add_heading(part_a_label or "Individualized Student Feedback", level=1)

        for student in part_a_results:
            _add_student_section(doc, student["name"], student["feedback"])

        doc.add_page_break()

    # ─── Part B: Collective Synthesis ─────────────────────────────
    if part_b_text:
        if has_part_a:
            _add_part_label(doc, "PART B")
        doc.add_heading(part_b_label, level=1)
        _add_markdown_content(doc, part_b_text)

    # Serialize to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
